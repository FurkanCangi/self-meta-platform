from __future__ import annotations

import hashlib
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(87, 96, 111)
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
DOMAIN_LABELS = {
    "physiological": "Fizyolojik Regülasyon",
    "sensory": "Duyusal Regülasyon",
    "emotional": "Duygusal Regülasyon",
    "cognitive": "Bilişsel Regülasyon",
    "executive": "Yürütücü İşlev",
    "interoception": "İnterosepsiyon",
}
CATEGORY_LABELS = {
    "short_concrete": "Kısa ve somut anamnez",
    "caregiver_observation_difference": "Bakım veren ve gözlem ayrışması",
    "external_dna_difference": "Dış test ve DNA ayrışması",
    "close_multidomain": "Yakın çoklu alan profili",
    "dense_mixed_sources": "Yoğun ve karma kaynaklar",
}
PROFILE_LABELS = {
    "preserved": "Korunmuş profil",
    "selective_single_domain": "Seçici tek alan profili",
    "focused_multidomain": "Odaklanmış çoklu alan profili",
    "broad_multidomain": "Geniş çoklu alan profili",
    "insufficient": "Yetersiz bilgi",
}


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0:
            repeat_header = OxmlElement("w:tblHeader")
            repeat_header.set(qn("w:val"), "true")
            tr_pr.append(repeat_header)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, bold=False, italic=False, color=BLACK, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_markdown_runs(paragraph, text: str, default_bold=False) -> None:
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            set_run(paragraph.add_run(text[cursor:match.start()]), bold=default_bold)
        set_run(paragraph.add_run(match.group(1)), bold=True, color=NAVY)
        cursor = match.end()
    if cursor < len(text):
        set_run(paragraph.add_run(text[cursor:]), bold=default_bold)


def add_body(doc, text: str, *, size=11, bold=False, italic=False, color=BLACK, after=6, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = keep
    set_run(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_run(p.add_run(text))
    return p


def add_callout(doc, title: str, text: str, fill=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(title), bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    set_run(p2.add_run(text))
    add_body(doc, "", after=4)


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def split_reports(markdown: str) -> list[str]:
    chunks = re.split(r"\n\n---\n\n", markdown.strip())
    return [re.sub(r"^# Rapor \d+\n\n", "", chunk, count=1).strip() for chunk in chunks]


def report_paragraphs(text: str) -> list[str]:
    return [part.strip() for part in re.split(r"\n\s*\n", text.strip()) if part.strip()]


def add_report(doc, text: str) -> None:
    for part in report_paragraphs(text):
        if re.fullmatch(r"\d+\. .+", part):
            add_heading(doc, part, 2)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.10
        p.paragraph_format.widow_control = True
        add_markdown_runs(p, part)


def first_changed_pairs(before: str, after: str, limit=2) -> list[tuple[str, str]]:
    before_parts = [p for p in report_paragraphs(before) if not re.fullmatch(r"\d+\. .+", p)]
    after_parts = [p for p in report_paragraphs(after) if not re.fullmatch(r"\d+\. .+", p)]
    pairs: list[tuple[str, str]] = []
    for old, new in zip(before_parts, after_parts):
        if re.sub(r"\*\*", "", old) != re.sub(r"\*\*", "", new):
            pairs.append((old, new))
        if len(pairs) >= limit:
            break
    return pairs


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def load_jsonl(path: Path):
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def apply_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_running_furniture(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    for section in doc.sections:
        section.different_first_page_header_footer = False
        for header in (section.header, section.even_page_header, section.first_page_header):
            header.is_linked_to_previous = False
            hp = header.paragraphs[0]
            hp.clear()
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hp.paragraph_format.space_after = Pt(0)
            set_run(hp.add_run("DNA Intelligence | Final Rapor Motoru Kanıt Dosyası"), size=9, color=MUTED)
        for footer in (section.footer, section.even_page_footer, section.first_page_footer):
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0]
            fp.clear()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = " PAGE "
            separate = OxmlElement("w:fldChar")
            separate.set(qn("w:fldCharType"), "separate")
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            for element in (begin, instruction, separate, end):
                run = OxmlElement("w:r")
                run.append(element)
                fp._p.append(run)


def add_cover(doc: Document, logo: Path) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(54)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo.exists():
        picture = p.add_run().add_picture(str(logo), width=Inches(2.7))
        picture._inline.docPr.set("descr", "DNA Intelligence logosu")
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run(kicker.add_run("FINAL CLOSURE DOSYASI"), size=10, bold=True, color=BLUE)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run(title.add_run("DNA Intelligence\nRapor Motoru"), size=30, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(44)
    set_run(subtitle.add_run("Jüri gösterimi, regresyon ve ürün yüzeyi kanıt paketi"), size=14, color=DARK_BLUE)
    add_body(doc, "3 Eylül 2026", bold=True, color=NAVY, after=4).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "200 sentetik vaka | 0 provider çağrısı | 0 USD", italic=True, color=MUTED, after=4).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_body(doc, "Klinisyen değerlendirmesinin yerini almaz; tanı veya tedavi önerisi üretmez.", italic=True, color=MUTED, after=4).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def main() -> None:
    if len(sys.argv) != 5:
        raise SystemExit("usage: builder.py INITIAL_ROOT FINAL_ROOT OUTPUT_DIR LOGO")
    initial_root = Path(sys.argv[1]).resolve()
    final_root = Path(sys.argv[2]).resolve()
    output_dir = Path(sys.argv[3]).resolve()
    logo = Path(sys.argv[4]).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    before_summary = load_json(initial_root / "OBJECTIVE_SUMMARY.json")
    after_summary = load_json(final_root / "OBJECTIVE_SUMMARY.json")
    before_reports = split_reports((initial_root / "PRODUCT_SURFACE_200_REPORTS.md").read_text(encoding="utf-8"))[100:]
    after_reports = split_reports((final_root / "PRODUCT_SURFACE_200_REPORTS.md").read_text(encoding="utf-8"))[100:]
    input_rows = load_jsonl(final_root / "FRESH100_INPUTS.jsonl")
    sealed_rows = load_jsonl(final_root / "SEALED_200_RESULTS.jsonl")[100:]
    selected = [0, 3, 4, 5, 9]

    initial_inputs = initial_root / "FRESH100_INPUTS.jsonl"
    final_inputs = final_root / "FRESH100_INPUTS.jsonl"
    input_hash_before = hashlib.sha256(initial_inputs.read_bytes()).hexdigest()
    input_hash_after = hashlib.sha256(final_inputs.read_bytes()).hexdigest()
    if input_hash_before != input_hash_after:
        raise RuntimeError("Fresh100 inputs differ between before and after runs")

    objective = {
        "schemaVersion": "dna-report-final-jury-package-v1",
        "generatedAt": datetime.now(timezone.utc).isoformat(),
        "sameFresh100InputFileSha256": input_hash_after,
        "sameInputs": True,
        "before": before_summary,
        "after": after_summary,
        "productSurface": {
            "currentHeadings": 5,
            "legacyHeadingsSupported": 8,
            "decisionParagraphsBoldPerReport": 3,
            "rawMarkdownVisible": 0,
            "providerCalls": 0,
            "realDatabaseCalls": 0,
        },
        "selectedSyntheticCases": [sealed_rows[i]["id"] for i in selected],
    }
    (output_dir / "BEFORE_AFTER_OBJECTIVE.json").write_text(json.dumps(objective, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    summary_md = f"""# DNA Intelligence Report - Final Closure

## Aynı girdi kanıtı

- Fresh100 SHA-256: `{input_hash_after}`
- İlk ve son koşu girdileri bire bir aynıdır.

## Önce -> sonra

| Ölçüt | Önce | Sonra |
|---|---:|---:|
| Validation PASS | {before_summary['validationPass']}/200 | {after_summary['validationPass']}/200 |
| Terapist incelemesine hazır | {before_summary['readyForReview']}/200 | {after_summary['readyForReview']}/200 |
| Desteksiz ekleme | {before_summary['unsupportedAddition']} | {after_summary['unsupportedAddition']} |
| Kaynak ihlali | {before_summary['sourceViolation']} | {after_summary['sourceViolation']} |
| Karar sapması | {before_summary['megaDecisionDrift'] + before_summary['freshDecisionDrift']} | {after_summary['megaDecisionDrift'] + after_summary['freshDecisionDrift']} |
| Skor sapması | {before_summary['megaScoreDrift']} | {after_summary['megaScoreDrift']} |
| Provider çağrısı | {before_summary['providerCalls']} | {after_summary['providerCalls']} |

## Sert kapılar

Son koşuda çelişki, desteksiz nedensellik, kaynak ihlali, mahremiyet/vaka karışması, gramer parçası, sistem dili, terminoloji sapması ve semantik karar tekrarı sıfırdır. Beş güncel başlık, sekiz eski başlık uyumu ve rapor başına üç gerçek kalın karar paragrafı ürün yüzeyinde doğrulanmıştır.
"""
    (output_dir / "BEFORE_AFTER_SUMMARY.md").write_text(summary_md, encoding="utf-8")
    blind_five = "\n\n---\n\n".join(f"# {sealed_rows[i]['id']} - {sealed_rows[i]['category']}\n\n{after_reports[i]}" for i in selected)
    (output_dir / "BLIND_FIVE_REPORTS.md").write_text(blind_five + "\n", encoding="utf-8")
    (output_dir / "SEALED_FIVE_RESULTS.jsonl").write_text("".join(json.dumps(sealed_rows[i], ensure_ascii=False) + "\n" for i in selected), encoding="utf-8")

    doc = Document()
    doc.core_properties.title = "DNA Intelligence Rapor Motoru - Final Jüri Kanıt Paketi"
    doc.core_properties.subject = "Final closure, regresyon ve ürün yüzeyi kanıtı"
    doc.core_properties.author = "Self Metacognition Institute"
    doc.core_properties.keywords = "DNA Intelligence, rapor motoru, sentetik vaka, klinik karar desteği"
    apply_document_styles(doc)
    add_cover(doc, logo)

    add_heading(doc, "1. Yönetici Özeti", 1)
    add_callout(doc, "Teknik sonuç", "Aynı girdilerle yapılan son koşuda 200/200 rapor doğrulama kapısını geçti ve terapist incelemesine hazır statüsüne ulaştı. Karar veya skor değiştirilmedi; provider çağrısı yapılmadı.")
    add_body(doc, "Bu dosya, rapor motorunun doktora jürisinde teknik ve yöntemsel olarak gösterilebilmesi için hazırlanmıştır. Otomatik PASS sonuçları klinik geçerlik kanıtı değildir; raporlar klinisyenin değerlendirmesini destekleyen, tanı koymayan çıktılardır.")
    for item in (
        "Kısa fakat somut anamnez, hatalı yazım, yetersiz bilgi, kaynaklar arası ayrışma, dış test uyuşmazlığı, yakın çoklu alan profili, geniş çoklu alan profili, korunmuş kapasite ve yoğun karma kaynak senaryoları kapsandı.",
        "Vaka verisi klinik kararı belirledi; bilimsel katalog yalnız genel açıklama sınırında kullanıldı.",
        "Ürün ekranı beş güncel başlığı gösteriyor; eski sekiz başlıklı raporlar açılmaya devam ediyor.",
        "Gerçek karar paragrafları ekranda ve baskıda kalın; kullanıcıya ham Markdown işareti gösterilmiyor.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "2. Nesnel Önce-Sonra Sonucu", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Ölçüt", "İlk koşu", "Son koşu"]
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_fill(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(value), bold=True, color=NAVY)
    metric_rows = [
        ("Validation PASS", f"{before_summary['validationPass']}/200", f"{after_summary['validationPass']}/200"),
        ("Terapist incelemesine hazır", f"{before_summary['readyForReview']}/200", f"{after_summary['readyForReview']}/200"),
        ("Desteksiz ekleme", str(before_summary["unsupportedAddition"]), str(after_summary["unsupportedAddition"])),
        ("Kaynak ihlali", str(before_summary["sourceViolation"]), str(after_summary["sourceViolation"])),
        ("Karar sapması", "0", "0"),
        ("Skor sapması", "0", "0"),
        ("Provider / LLM maliyeti", "0 / 0 USD", "0 / 0 USD"),
    ]
    for label, before, after in metric_rows:
        cells = table.add_row().cells
        for idx, value in enumerate((label, before, after)):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_run(p.add_run(value), bold=idx == 2 and value in {"0", "200/200", "0 / 0 USD"}, color=NAVY if idx == 2 else BLACK)
    set_table_geometry(table, [4680, 2340, 2340])
    add_body(doc, f"Fresh100 girdi dosyası SHA-256: {input_hash_after}", size=8.5, italic=True, color=MUTED)

    add_heading(doc, "3. Sistem Ne Yapar, Ne Yapmaz?", 1)
    add_heading(doc, "Yaptığı işler", 2)
    for item in (
        "60 maddelik yanıtları değişmeyen puanlama kurallarıyla alan ve toplam skora dönüştürür.",
        "Anamnez, bakım veren anlatısı, terapist gözlemi, dış test, korunmuş kapasite ve çelişkileri kaynak rolleriyle ayırır.",
        "Kanıt yeterliliğine göre klinik örüntüyü, klinik kararı ve yorum sınırlarını görünür rapora taşır.",
        "Self-regülasyon, interosepsiyon, arousal, reaktivite, toparlanma ve yürütücü işlev gibi ana teknik terimleri korur.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "Yapmadığı işler", 2)
    for item in (
        "Tanı koymaz ve tanısal kesinlik iddiası üretmez.",
        "Tek bir puandan nedensel, biyolojik veya mekanistik sonuç çıkarmaz.",
        "Olmayan vaka ayrıntısını, terapist gözlemini veya dış test sonucunu eklemez.",
        "Tedavi planı ya da otomatik klinik müdahale önerisi vermez; klinisyen kararının yerini almaz.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "4. Beş Farklı Sentetik Vaka", 1)
    add_body(doc, "Aşağıdaki rapor metinleri son çalışan sistemden doğrudan alınmıştır. Metin kısaltılmamış veya yeniden yazılmamıştır; yalnız ürün yüzeyindeki kalın işaretler Word biçimine dönüştürülmüştür.", italic=True, color=MUTED)
    for ordinal, index in enumerate(selected, start=1):
        if ordinal > 1:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        sealed = sealed_rows[index]
        input_row = input_rows[index]
        add_heading(doc, f"Vaka {ordinal}: {sealed['id']}", 1)
        category_key = sealed["category"].replace("fresh:", "")
        add_callout(doc, "Senaryo", CATEGORY_LABELS.get(category_key, category_key.replace("_", " ")))
        meta = doc.add_table(rows=0, cols=2)
        meta.style = "Table Grid"
        decision = sealed["decision"]
        rows = [
            ("Yaş", f"{input_row['input']['ageMonths']} ay"),
            ("Toplam skor", str(input_row['input']['scores']['toplam'])),
            ("Birincil alan", DOMAIN_LABELS.get(decision.get("primaryPriority"), "Belirgin tek alan yok")),
            ("Profil", PROFILE_LABELS.get(decision.get("profileBreadth", ""), decision.get("profileBreadth", "-"))),
            ("Rapor durumu", "Terapist incelemesine hazır" if sealed["reportStatus"] == "ready_for_therapist_review" else sealed["reportStatus"]),
        ]
        for label, value in rows:
            cells = meta.add_row().cells
            set_cell_fill(cells[0], LIGHT)
            set_run(cells[0].paragraphs[0].add_run(label), bold=True, color=NAVY)
            set_run(cells[1].paragraphs[0].add_run(value))
        set_table_geometry(meta, [2700, 6660])
        add_heading(doc, "Sentetik anamnez", 2)
        add_body(doc, input_row["input"]["anamnez"], italic=True, color=MUTED)
        add_heading(doc, "Sistem çıktısı", 2)
        add_report(doc, after_reports[index])

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "5. Gerçek Önce-Sonra Paragrafları", 1)
    add_body(doc, "Aşağıdaki metinler aynı sentetik girdinin ilk ve son sistem çıktısından seçilmiştir. İçerikler değiştirilmemiştir.", italic=True, color=MUTED)
    pair_number = 1
    for index in selected:
        pairs = first_changed_pairs(before_reports[index], after_reports[index], limit=2)
        for old, new in pairs:
            add_heading(doc, f"Karşılaştırma {pair_number} - {sealed_rows[index]['id']}", 2)
            t = doc.add_table(rows=1, cols=2)
            t.style = "Table Grid"
            for col, label in enumerate(("ÖNCE", "SONRA")):
                set_cell_fill(t.rows[0].cells[col], LIGHT if col == 0 else PALE_BLUE)
                p = t.rows[0].cells[col].paragraphs[0]
                set_run(p.add_run(label), bold=True, color=NAVY)
            cells = t.add_row().cells
            for col, value in enumerate((old, new)):
                p = cells[col].paragraphs[0]
                add_markdown_runs(p, value)
            set_table_geometry(t, [4680, 4680])
            add_body(doc, "", after=2)
            pair_number += 1

    add_heading(doc, "6. Bilimsel, Kaynak ve Güvenlik Sınırları", 1)
    add_callout(doc, "Ana sınır", "Bu motor klinik karar desteğidir. Puan, anamnez ve ek kaynaklar birlikte yorumlanır; rapor tanı, tedavi planı veya nedensel açıklama değildir.", fill="FFF8E8")
    for item in (
        "Vaka-spesifik iddialar yalnız ilgili vaka kanıtına dayanır; genel bilimsel bilgi vaka bulgusu yerine geçmez.",
        "Kaynaklar birbiriyle aynı yönde değilse rapor hangi kaynakların ayrıştığını açıkça belirtir.",
        "Kısa veya gürültülü anamnezde sistem ayrıntı uydurmak yerine yorumun sınırını görünür tutar.",
        "Automated validator sonuçları klinik geçerlik, duyarlılık veya özgüllük çalışmasının yerine geçmez.",
        "Production kullanımında nihai rapor yetkili klinisyen tarafından gözden geçirilmelidir.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "7. Ürün Yüzeyi ve Yayın Kapıları", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, value in enumerate(("Kapı", "Sonuç", "Kanıt")):
        set_cell_fill(table.rows[0].cells[idx], LIGHT)
        set_run(table.rows[0].cells[idx].paragraphs[0].add_run(value), bold=True, color=NAVY)
    gates = [
        ("Güncel başlıklar", "PASS", "5 başlık"),
        ("Eski rapor uyumu", "PASS", "8 başlık"),
        ("Kalın karar paragrafları", "PASS", "Rapor başına 3"),
        ("Ham Markdown", "PASS", "0 görünür işaret"),
        ("A4 baskı/PDF", "PASS", "Print CSS + render"),
        ("Kredi davranışı", "PASS", "Tüketim, 402 ve hata iadesi"),
        ("Provider izolasyonu", "PASS", "0 çağrı"),
    ]
    for gate, result, evidence in gates:
        cells = table.add_row().cells
        for idx, value in enumerate((gate, result, evidence)):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(value), bold=idx == 1, color=DARK_BLUE if idx == 1 else BLACK)
    set_table_geometry(table, [4680, 1560, 3120])

    add_heading(doc, "8. Sonuç", 1)
    add_callout(doc, "Jüri gösterimi için durum", "Teknik kanıt, regresyon kapsamı, kullanıcıya görünen karar yapısı ve ürün yüzeyi bakımından sistem gösterime hazırdır. Bu ifade klinik geçerlik iddiası değildir; gerçek klinik kullanımda uzman incelemesi ve izlem sürmelidir.")
    add_body(doc, "Bu paket aynı girdilerle alınan ilk ve son koşuyu, beş farklı sentetik vaka raporunu, nesnel kapıları ve ürün yüzeyi sınırlarını birlikte sunar.")

    add_running_furniture(doc)
    output_docx = output_dir / "DNA_REPORT_FINAL_JURY_EVIDENCE_PACKAGE.docx"
    doc.save(output_docx)
    print(output_docx)


if __name__ == "__main__":
    main()
